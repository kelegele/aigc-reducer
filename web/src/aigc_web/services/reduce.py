# web/src/aigc_web/services/reduce.py
"""P3 检测/改写业务逻辑。"""

import asyncio
import io
import json
import logging
import threading
import time
from typing import AsyncGenerator

from sqlalchemy.orm import Session

from aigc_reducer_core import CancelledError
from aigc_reducer_core.detector import AIGCDetector
from aigc_reducer_core.llm_client import LLMClient
from aigc_reducer_core.parser import Paragraph
from aigc_reducer_core.rewriter import Rewriter

from aigc_web.config import settings
from aigc_web.models.reduction_paragraph import ReductionParagraph
from aigc_web.models.reduction_task import ReductionTask
from aigc_web.services import credit as credit_service

logger = logging.getLogger(__name__)


class ReduceService:
    """改写任务服务。"""

    def __init__(self, db: Session):
        self.db = db
        self._llm_client = LLMClient(
            model=settings.LLM_MODEL,
            api_key=settings.LLM_API_KEY,
            base_url=settings.LLM_BASE_URL,
        )

    # ── 任务创建 ──

    async def create_task(
        self,
        user_id: int,
        detect_mode: str,
        style: str,
        text: str | None = None,
        file_path: str | None = None,
    ) -> ReductionTask:
        """创建改写任务。解析文本/文件为段落。"""
        # 1. 获取原文并解析段落
        if text:
            original_text = text
            paragraphs = []
            for i, block in enumerate(text.split("\n\n")):
                block = block.strip()
                if not block:
                    continue
                # 识别 markdown 标题
                if block.startswith("#"):
                    heading_text = block.lstrip("#").strip()
                    if heading_text:
                        paragraphs.append(Paragraph(text=heading_text, index=i, is_heading=True))
                    continue
                paragraphs.append(Paragraph(text=block, index=i))
        elif file_path:
            paragraphs = await asyncio.to_thread(
                _parse_document_sync, file_path
            )
            original_text = "\n\n".join(p.text for p in paragraphs)
        else:
            raise ValueError("必须提供 text 或 file_path")

        # 2. 创建任务记录
        first_heading = next((p.text for p in paragraphs if p.is_heading), None)
        if first_heading:
            title = first_heading[:100]
        else:
            # 无标题段落时取第一行（避免混入正文内容）
            first_line = original_text.split("\n")[0].strip()
            if len(first_line) > 100:
                title = first_line[:100] + "..."
            else:
                title = first_line
        task = ReductionTask(
            user_id=user_id,
            title=title,
            status="parsing",
            detect_mode=detect_mode,
            style=style,
            original_text=original_text,
        )
        self.db.add(task)
        self.db.flush()

        # 3. 创建段落记录
        for p in paragraphs:
            para = ReductionParagraph(
                task_id=task.id,
                index=p.index,
                original_text=p.text,
                is_heading=p.is_heading,
                has_formula=p.has_formula,
                has_code=p.has_code,
            )
            self.db.add(para)

        task.status = "detecting"
        self.db.commit()
        self.db.refresh(task)
        logger.info(
            "[create_task] task=%s user=%d mode=%s style=%s paragraphs=%d",
            task.id[:8], user_id, detect_mode, style, len(paragraphs),
        )
        return task

    # ── 检测（SSE） ──

    async def start_detection(self, task_id: str) -> AsyncGenerator[dict, None]:
        """启动检测，SSE 生成器。detecting 状态视为上次中断，重置后重新检测。"""
        task = self._get_task(task_id)
        if task.status not in ("detecting", "detected"):
            logger.warning("[detect] task=%s invalid status=%s", task_id[:8], task.status)
            yield {"type": "error", "message": f"任务状态不正确: {task.status}"}
            return

        logger.info("[detect] task=%s start, mode=%s", task_id[:8], task.detect_mode)
        t0 = time.time()

        # detecting 状态 = 上次中断，重置段落检测结果
        if task.status == "detecting":
            paragraphs = self._get_paragraphs(task_id)
            for p in paragraphs:
                p.detection_result = None
                p.risk_level = None
                p.needs_processing = False
                p.status = "pending"
            self.db.flush()
            logger.info("[detect] task=%s reset interrupted paragraphs", task_id[:8])

        paragraphs = self._get_paragraphs(task_id)
        parsed = [
            Paragraph(
                text=p.original_text,
                index=p.index,
                is_heading=p.is_heading,
                has_formula=p.has_formula,
                has_code=p.has_code,
            )
            for p in paragraphs
        ]

        # LLM 模式需要预检积分
        if task.detect_mode == "llm":
            est = self._estimate_tokens(parsed)
            cost = self._tokens_to_credits(est)
            balance = credit_service.get_balance(self.db, task.user_id)
            if balance < cost:
                logger.warning("[detect] task=%s insufficient balance=%d cost=%d", task_id[:8], balance, cost)
                yield {
                    "type": "error",
                    "message": f"积分不足，预计需要 {cost} 积分",
                }
                return

        # 执行检测
        cancel_event = threading.Event()
        detector = AIGCDetector(
            mode=task.detect_mode,
            cancel_event=cancel_event,
            llm_client=self._llm_client if task.detect_mode == "llm" else None,
        )

        try:
            results = await asyncio.to_thread(detector.analyze_all, parsed)
        except CancelledError:
            task.status = "failed"
            self.db.commit()
            logger.info("[detect] task=%s cancelled", task_id[:8])
            yield {"type": "error", "message": "检测已取消"}
            return
        except Exception as e:
            task.status = "failed"
            self.db.commit()
            logger.error("[detect] task=%s failed: %s", task_id[:8], e, exc_info=True)
            yield {"type": "error", "message": f"检测失败: {e}"}
            return

        # 保存检测结果
        total = len(results)
        needs_processing_count = 0
        for i, result in enumerate(results):
            para = paragraphs[i]
            para.detection_result = result
            para.risk_level = result.get("risk_level")
            para.needs_processing = result.get("risk_level") != "low" and not para.is_heading
            para.status = "detected"
            if para.needs_processing:
                needs_processing_count += 1

            yield {
                "type": "paragraph_done",
                "index": i,
                "risk_level": result.get("risk_level", ""),
                "composite_score": result.get("composite_score", 0),
                "current": i + 1,
                "total": total,
            }

        # LLM 模式扣费
        if task.detect_mode == "llm":
            est = self._estimate_tokens(parsed)
            cost = self._tokens_to_credits(est)
            credit_service.consume(
                self.db,
                task.user_id,
                est,
                ref_type="reduction_task",
                ref_id=task_id,
                remark="LLM 检测",
            )
            task.total_tokens += est
            task.total_credits += cost
            logger.info("[detect] task=%s LLM billed: tokens=%d credits=%d", task_id[:8], est, cost)

        try:
            task.status = "detected"
            self.db.commit()
        except Exception as e:
            self.db.rollback()
            task.status = "failed"
            self.db.commit()
            logger.error("[detect] task=%s commit failed: %s", task_id[:8], e, exc_info=True)
            yield {"type": "error", "message": f"保存检测结果失败: {e}"}
            return

        elapsed = time.time() - t0
        logger.info(
            "[detect] task=%s done: total=%d needs_processing=%d elapsed=%.1fs",
            task_id[:8], total, needs_processing_count, elapsed,
        )

        yield {
            "type": "complete",
            "total_paragraphs": total,
            "needs_processing": needs_processing_count,
        }

    # ── 全量语义重构（SSE） ──

    async def start_reconstruction(self, task_id: str) -> AsyncGenerator[dict, None]:
        """全量语义重构——逐段重构并推送进度。"""
        task = self._get_task(task_id)
        if task.status != "detected":
            logger.warning("[reconstruct] task=%s invalid status=%s", task_id[:8], task.status)
            yield {"type": "error", "message": f"任务状态不正确: {task.status}"}
            return

        logger.info("[reconstruct] task=%s start", task_id[:8])
        t0 = time.time()

        paragraphs = self._get_paragraphs(task_id)
        body_paragraphs = [p for p in paragraphs if not p.is_heading]
        if not body_paragraphs:
            yield {"type": "error", "message": "没有可重构的段落"}
            return

        # 预检积分
        total_chars = sum(len(p.original_text) for p in body_paragraphs)
        est = max(1, int(total_chars * 1.5))
        cost = self._tokens_to_credits(est)
        balance = credit_service.get_balance(self.db, task.user_id)
        if balance < cost:
            logger.warning("[reconstruct] task=%s insufficient balance=%d cost=%d", task_id[:8], balance, cost)
            yield {
                "type": "error",
                "message": f"积分不足，预计需要 {cost} 积分",
            }
            return

        cancel_event = threading.Event()
        rewriter = Rewriter(
            task.style,
            llm_client=self._llm_client,
            cancel_event=cancel_event,
        )
        task.full_reconstruct = True

        total_est = 0
        total_cost = 0
        done = 0
        total = len(body_paragraphs)
        failed = False

        for para in body_paragraphs:
            if cancel_event.is_set():
                yield {"type": "error", "message": "重构已取消"}
                return

            try:
                rewritten = await asyncio.to_thread(
                    rewriter.rewrite_single,
                    para.original_text,
                    {"composite_score": 50},
                    conservative=False,
                )
            except CancelledError:
                logger.info("[reconstruct] task=%s cancelled at para %d", task_id[:8], para.index)
                yield {"type": "error", "message": "重构已取消"}
                return
            except Exception as e:
                logger.error("[reconstruct] task=%s para %d failed: %s", task_id[:8], para.index, e)
                yield {"type": "error", "message": f"重构段落 {para.index + 1} 失败: {e}"}
                failed = True
                break

            # 更新段落文本
            para.original_text = rewritten.strip()
            para_est = max(1, int(len(para.original_text) * 1.5))
            total_est += para_est
            total_cost += self._tokens_to_credits(para_est)
            done += 1

            yield {
                "type": "progress",
                "index": para.index,
                "current": done,
                "total": total,
            }

        if failed:
            task.status = "failed"
            self.db.commit()
            return

        # 更新全文
        task.original_text = "\n\n".join(p.original_text for p in paragraphs)

        # 扣费
        credit_service.consume(
            self.db,
            task.user_id,
            total_est,
            ref_type="reduction_task",
            ref_id=task_id,
            remark="全量重构",
        )
        task.total_tokens += total_est
        task.total_credits += total_cost

        self.db.commit()
        elapsed = time.time() - t0
        logger.info("[reconstruct] task=%s done: %d paras, credits=%d elapsed=%.1fs", task_id[:8], done, total_cost, elapsed)
        yield {"type": "complete", "credits_used": total_cost}

    # ── 改写（SSE） ──

    async def start_rewrite(self, task_id: str) -> AsyncGenerator[dict, None]:
        """启动改写，SSE 生成器。对 needs_processing 的段落生成 A/B 选项。"""
        task = self._get_task(task_id)
        if task.status not in ("detected", "rewriting"):
            logger.warning("[rewrite] task=%s invalid status=%s", task_id[:8], task.status)
            yield {"type": "error", "message": f"任务状态不正确: {task.status}"}
            return

        task.status = "rewriting"
        self.db.flush()

        paragraphs = [p for p in self._get_paragraphs(task_id) if p.needs_processing and not p.is_heading]
        if not paragraphs:
            task.status = "rewritten"
            self.db.commit()
            logger.info("[rewrite] task=%s no paragraphs need processing, done", task_id[:8])
            yield {"type": "complete", "total_credits_used": 0}
            return

        logger.info("[rewrite] task=%s start: %d paragraphs to rewrite", task_id[:8], len(paragraphs))
        t0 = time.time()

        cancel_event = threading.Event()
        rewriter = Rewriter(
            task.style,
            llm_client=self._llm_client,
            cancel_event=cancel_event,
        )

        total = len(paragraphs)
        total_credits = 0

        for i, para in enumerate(paragraphs):
            # 逐段预检积分
            est = self._estimate_tokens_for_text(para.original_text) * 2  # A+B 两次调用
            cost = self._tokens_to_credits(est)
            balance = credit_service.get_balance(self.db, task.user_id)
            if balance < cost:
                task.status = "failed"
                self.db.commit()
                logger.warning(
                    "[rewrite] task=%s insufficient balance at para %d: balance=%d cost=%d",
                    task_id[:8], i, balance, cost,
                )
                yield {
                    "type": "error",
                    "message": f"积分不足，段落 {i + 1} 预计需要 {cost} 积分",
                }
                return

            yield {"type": "progress", "current": i + 1, "total": total}

            detection = para.detection_result or {}
            para_t0 = time.time()

            try:
                aggressive = await asyncio.to_thread(
                    rewriter.rewrite_single,
                    para.original_text,
                    detection,
                    conservative=False,
                )
                conservative = await asyncio.to_thread(
                    rewriter.rewrite_single,
                    para.original_text,
                    detection,
                    conservative=True,
                )
            except CancelledError:
                task.status = "failed"
                self.db.commit()
                logger.info("[rewrite] task=%s cancelled at para %d", task_id[:8], i)
                yield {"type": "error", "message": "改写已取消"}
                return
            except Exception as e:
                task.status = "failed"
                self.db.commit()
                logger.error(
                    "[rewrite] task=%s para %d failed: %s", task_id[:8], i, e, exc_info=True,
                )
                yield {"type": "error", "message": f"改写失败: {e}"}
                return

            para.rewrite_aggressive = aggressive
            para.rewrite_conservative = conservative
            para.status = "rewritten"

            # 扣费
            credit_service.consume(
                self.db,
                task.user_id,
                est,
                ref_type="reduction_task",
                ref_id=task_id,
                remark=f"段落 {para.index} 改写",
            )
            total_credits += cost
            task.total_tokens += est
            task.total_credits += cost

            para_elapsed = time.time() - para_t0
            logger.info(
                "[rewrite] task=%s para %d/%d done: index=%d est_tokens=%d credits=%d elapsed=%.1fs",
                task_id[:8], i + 1, total, para.index, est, cost, para_elapsed,
            )

            yield {
                "type": "paragraph_ready",
                "index": para.index,
                "aggressive": aggressive,
                "conservative": conservative,
            }

        task.status = "rewritten"
        self.db.commit()
        elapsed = time.time() - t0
        logger.info(
            "[rewrite] task=%s all done: paragraphs=%d total_credits=%d elapsed=%.1fs",
            task_id[:8], total, total_credits, elapsed,
        )
        yield {"type": "complete", "total_credits_used": total_credits}

    # ── 段落确认 ──

    def confirm_paragraph(
        self,
        task_id: str,
        index: int,
        choice: str,
        manual_text: str | None = None,
    ) -> ReductionParagraph:
        """确认段落选择。"""
        para = (
            self.db.query(ReductionParagraph)
            .filter_by(task_id=task_id, index=index)
            .first()
        )
        if not para:
            raise ValueError(f"段落不存在: {index}")

        para.user_choice = choice
        if choice == "aggressive":
            para.final_text = para.rewrite_aggressive
        elif choice == "conservative":
            para.final_text = para.rewrite_conservative
        elif choice == "original":
            para.final_text = para.original_text
        elif choice == "manual":
            para.final_text = manual_text or ""
        else:
            raise ValueError(f"无效选择: {choice}")

        para.status = "confirmed"
        self.db.commit()
        logger.info("[confirm] task=%s para=%d choice=%s", task_id[:8], index, choice)
        return para

    # ── 任务完成 ──

    def finalize_task(self, task_id: str) -> ReductionTask:
        """所有段落确认后生成最终文档。"""
        task = self._get_task(task_id)
        paragraphs = self._get_paragraphs(task_id)

        for p in paragraphs:
            if p.status != "confirmed":
                if not p.needs_processing:
                    # 低风险段落自动用原文
                    p.final_text = p.original_text
                    p.status = "confirmed"
                    p.user_choice = "original"
                else:
                    raise ValueError(f"段落 {p.index} 尚未确认")

        # 拼接最终文本
        sorted_paras = sorted(paragraphs, key=lambda p: p.index)
        task.reduced_text = "\n\n".join(
            p.final_text for p in sorted_paras if p.final_text
        )
        task.status = "completed"
        self.db.commit()
        self.db.refresh(task)
        logger.info("[finalize] task=%s done: %d paragraphs", task_id[:8], len(paragraphs))
        return task

    # ── 查询 ──

    def get_task(self, task_id: str, user_id: int) -> ReductionTask:
        """获取用户的任务。"""
        task = (
            self.db.query(ReductionTask)
            .filter_by(id=task_id, user_id=user_id)
            .first()
        )
        if not task:
            raise ValueError("任务不存在")
        return task

    def list_tasks(
        self,
        user_id: int,
        page: int = 1,
        page_size: int = 10,
        status: str | None = None,
        keyword: str | None = None,
    ) -> tuple[list[ReductionTask], int]:
        """分页查询用户的改写任务，支持状态筛选和标题搜索。"""
        query = (
            self.db.query(ReductionTask)
            .filter_by(user_id=user_id)
            .order_by(ReductionTask.created_at.desc())
        )
        if status == "in_progress":
            query = query.filter(
                ReductionTask.status.notin_(["completed", "failed"])
            )
        elif status:
            query = query.filter_by(status=status)
        if keyword:
            query = query.filter(
                ReductionTask.title.ilike(f"%{keyword}%")
                | ReductionTask.id.ilike(f"%{keyword}%")
            )
        total = query.count()
        tasks = query.offset((page - 1) * page_size).limit(page_size).all()
        return tasks, total

    def estimate_credits(self, task_id: str, operation: str) -> dict:
        """预估积分消耗。"""
        task = self._get_task(task_id)
        paragraphs = self._get_paragraphs(task_id)

        if operation == "detect" and task.detect_mode == "llm":
            parsed = [
                Paragraph(text=p.original_text, index=p.index) for p in paragraphs
            ]
            est = self._estimate_tokens(parsed)
        elif operation == "reconstruct":
            total_chars = sum(len(p.original_text) for p in paragraphs)
            est = max(1, int(total_chars * 1.5))
        elif operation == "rewrite":
            needs = [p for p in paragraphs if p.needs_processing]
            est = sum(
                self._estimate_tokens_for_text(p.original_text) * 2 for p in needs
            )
        else:
            est = 0

        cost = self._tokens_to_credits(est)
        balance = credit_service.get_balance(self.db, task.user_id)

        return {
            "estimated_credits": cost,
            "current_balance": balance,
            "sufficient": balance >= cost,
        }

    # ── 私有方法 ──

    def _get_task(self, task_id: str) -> ReductionTask:
        task = self.db.query(ReductionTask).filter_by(id=task_id).first()
        if not task:
            raise ValueError("任务不存在")
        return task

    def _get_paragraphs(self, task_id: str) -> list[ReductionParagraph]:
        return (
            self.db.query(ReductionParagraph)
            .filter_by(task_id=task_id)
            .order_by(ReductionParagraph.index)
            .all()
        )

    @staticmethod
    def _estimate_tokens(paragraphs: list[Paragraph]) -> int:
        total_chars = sum(len(p.text) for p in paragraphs)
        return max(1, int(total_chars * 1.5))

    @staticmethod
    def _estimate_tokens_for_text(text: str) -> int:
        return max(1, int(len(text) * 1.5))

    @staticmethod
    def _tokens_to_credits(token_count: int) -> int:
        return max(1, int(token_count / 1000 * settings.CREDITS_PER_1K_TOKENS))


def _parse_document_sync(file_path: str) -> list[Paragraph]:
    """同步调用 parse_document，供 asyncio.to_thread 使用。"""
    from aigc_reducer_core.parser import parse_document

    return parse_document(file_path)


def _strip_markdown(text: str) -> str:
    """清除文本中的 Markdown 格式符号，返回纯文本。"""
    import re
    # 去除 # 标题标记
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    # 去除 **粗体** 和 __粗体__
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    # 去除 *斜体* 和 _斜体_
    text = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', text)
    text = re.sub(r'(?<!_)_(?!_)(.+?)(?<!_)_(?!_)', r'\1', text)
    # 去除 ~~删除线~~
    text = re.sub(r'~~(.+?)~~', r'\1', text)
    # 去除 `行内代码`
    text = re.sub(r'`(.+?)`', r'\1', text)
    # 去除 ```代码块```
    text = re.sub(r'```\w*\n?', '', text)
    # 去除 [链接文字](url)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    # 去除 ![图片](url)
    text = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', r'\1', text)
    # 去除 > 引用标记
    text = re.sub(r'^>\s?', '', text, flags=re.MULTILINE)
    # 去除 ---/___ 分割线
    text = re.sub(r'^[-_*]{3,}\s*$', '', text, flags=re.MULTILINE)
    return text.strip()


def export_docx(task: ReductionTask) -> io.BytesIO:
    """将改写结果导出为 DOCX 文件。"""
    from docx import Document

    doc = Document()
    paragraphs = sorted(task.paragraphs, key=lambda p: p.index)
    for p in paragraphs:
        text = _strip_markdown(p.final_text or p.original_text)
        if not text:
            continue
        if p.is_heading:
            doc.add_heading(text, level=2)
        else:
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    logger.info("[export] task=%s exported as docx", task.id[:8])
    return buf
